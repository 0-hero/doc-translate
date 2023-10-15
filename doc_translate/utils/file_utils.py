import logging
import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from .table_process import HtmlToDocx


def parse_layout(layout):
    logging.info("Parsing layout...")
    document_structure = {
        "document": {
            "pages": []
        }
    }
    for i, page in enumerate(layout.pages):
        page_structure = {
            "page_number": i,
            "layout_elements": []
        }

        # Process each LayoutElement
        for element in page.elements:
            element_structure = {
                "type": element.type,
                "text": element.text,
                "bbox": {
                    "x1": element.bbox.x1,
                    "y1": element.bbox.y1,
                    "x2": element.bbox.x2,
                    "y2": element.bbox.y2
                },
                "source": str(element.source),  # Assuming source is an enum and needs to be converted to string
                "prob": element.prob,
                "image_path": element.image_path,
                "parent": element.parent  # Assuming this is a reference to another LayoutElement or None
            }
            page_structure["layout_elements"].append(element_structure)

        document_structure["document"]["pages"].append(page_structure)
    logging.info("Layout parsed successfully.")
    return document_structure


def convert_info_docx(document_structure, save_folder, docx_name):
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(10)  # Adjusting the default font size

    for page in document_structure["document"]["pages"]:
        layout_elements = page["layout_elements"]

        for element in layout_elements:
            element_type = element['type'].lower()  # Normalize the type value
            text = element.get('translated_text', '')  # Default to an empty string if 'text' is not provided

            if element_type == 'title':
                doc.add_heading(text, level=1)

            elif element_type == 'headline':
                doc.add_heading(text, level=2)

            elif element_type == 'subheadline':
                doc.add_heading(text, level=3)

            elif element_type == 'abstract':
                paragraph = doc.add_paragraph()
                text_run = paragraph.add_run(text)
                text_run.italic = True  # Stylistic choice to italicize abstracts

            elif element_type in ['address', 'author']:
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Addresses and authors are often right-aligned
                paragraph.add_run(text)

            elif element_type == 'caption':
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Captions are often centered
                paragraph.add_run(text)

            elif element_type in ['footer', 'header']:
                # Special handling if docx library supports header and footer
                paragraph = doc.add_paragraph()
                paragraph.add_run("[{}] ".format(element_type.title()) + text)  # Adding a note for clarity

            elif element_type == 'metadata':
                # Handling depends on the nature of the metadata
                paragraph = doc.add_paragraph()
                paragraph.add_run("[Metadata] " + text)  # Adding a note for clarity

            elif element_type == 'page number':
                # Special handling if docx library supports footer functionality for page numbers
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Page numbers are often right-aligned
                paragraph.add_run(text)

            elif element_type == 'value':
                # Handling depends on the context of 'value'
                paragraph = doc.add_paragraph()
                paragraph.add_run(text)

            elif element_type in ['advertisement', 'field-name', 'formula', 'link', 'misc']:
                paragraph = doc.add_paragraph()
                text_run = paragraph.add_run(text)
                if element_type == 'link':
                    text_run.font.color.rgb = RGBColor(0, 0, 255)  # Making links blue
                    text_run.underline = True

            elif element_type in ['list', 'list-item']:
                paragraph = doc.add_paragraph(style='ListBullet' if element_type == 'list' else 'ListNumber')
                paragraph.add_run(text)

            elif element_type in ['picture', 'chart']:
                paragraph = doc.add_paragraph()
                paragraph.add_run("***Picture Placeholder***")
                # img_path = element.get('image_path')
                # if img_path and os.path.exists(img_path):
                #     paragraph_pic = doc.add_paragraph()
                #     paragraph_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #     run = paragraph_pic.add_run()
                #     run.add_picture(img_path, width=Inches(4))  # Adjust image width as needed

            elif element_type == 'table':
                # Placeholder for table handling
                # Implement table addition based on your specific table data structure
                paragraph = doc.add_paragraph()
                paragraph.add_run("***Table Placeholder***")

            elif element_type == 'text':
                paragraph = doc.add_paragraph()
                paragraph.add_run(text)

            elif element_type == 'threading':
                paragraph = doc.add_paragraph()
                paragraph.add_run(text)

            # Additional types can be added with elif statements and handled specifically

    # Save to docx
    docx_path = os.path.join(save_folder, '{}.docx'.format(docx_name))
    doc.save(docx_path)
    logging.info('docx saved to {}'.format(docx_path))